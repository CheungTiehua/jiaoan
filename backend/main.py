"""LeKai教案系统(K9-AI版) — v1.0"""

import io
import logging

_log = logging.getLogger("lekai")
import json as _json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from fastapi import FastAPI, HTTPException, Header, Depends, Request
from fastapi.responses import StreamingResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from rag import (
    generate_lesson, revise_lesson, generate_unit_plan, generate_reflection,
    generate_exam_analysis_bundle, generate_peer_analysis_bundle,
    generate_teaching_guide_bundle, prepare_lesson_evidence, stream_lesson_plan,
)
from teaching_evidence import (
    DOC_TYPE_LABELS, METHOD_DOC_TYPES, NORMATIVE_DOC_TYPES,
    build_generated_blocks, evidence_lookup,
    get_teaching_evidence_by_id, search_teaching_evidence,
    split_evidence, validate_block_citations,
)
from scripts.all_textbooks import GRADE_TEXTBOOKS
from auth import (
    register_user, login_user, logout_user, get_user_from_token,
    list_users, save_history, get_history, get_history_detail,
    get_user_role, save_history_mindmap, update_history_fields, audit_log,
)
from admin_api import (
    submit_for_review, get_review_queue, get_review_detail,
    approve_review, reject_review, get_dashboard_stats, set_user_role,
)
from collab import (
    create_group, list_groups, join_group, get_user_groups,
    share_plan, get_group_plans, assign_task, get_group_tasks,
    complete_task, add_comment,
)
from config import (
    VERSION, PROMPTS_FILE, LEKAI_ACCEPTANCE_MODE,
    PDF_OCR_DPI_SCALE, PDF_OCR_LANG, PDF_OCR_MAX_PAGES, PDF_UPLOAD_MAX_MB,
)
from security import check_rate_limit
from health import get_health
from backup import create_backup, restore_backup
from feedback import submit_feedback, get_feedback, get_stats as get_feedback_stats
from annotation import add_annotation, get_annotations, get_stats as get_annotation_stats
from mindmap import generate_dual_mindmap
from fastapi import UploadFile, File as FastAPIFile, Form

app = FastAPI(
    title="LeKai教案知识库 API",
    description="小学语文教案智能生成平台 — LeKai K9-AI版",
    version=VERSION
)

import os as _os
_cors_origins = [o.strip() for o in _os.getenv("LEKAI_CORS_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(",")]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins, allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

_source_pages_dir = Path(__file__).resolve().parent.parent / "data" / "source_pages"
_source_pages_dir.mkdir(parents=True, exist_ok=True)
app.mount("/source-pages", StaticFiles(directory=str(_source_pages_dir)), name="source-pages")
app.mount("/api/source-pages", StaticFiles(directory=str(_source_pages_dir)), name="api-source-pages")


# ---- Auth dependency ----

def require_auth(authorization: str = Header(default="")) -> str:
    """从 Authorization: Bearer <token> 中提取用户"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="请先登录")
    token = authorization[len("Bearer "):]
    username = get_user_from_token(token)
    if not username:
        raise HTTPException(status_code=401, detail="登录已过期，请重新登录")
    return username


def require_admin_or_reviewer(username: str = Depends(require_auth)) -> str:
    role = get_user_role(username)
    if role not in ("admin", "reviewer"):
        raise HTTPException(status_code=403, detail="需要管理员或教研组长权限")
    return username


def require_admin(username: str = Depends(require_auth)) -> str:
    if get_user_role(username) != "admin":
        raise HTTPException(status_code=403, detail="需要管理员权限")
    return username


# ---- Models ----

class AuthRequest(BaseModel):
    username: str
    password: str


class GenerateRequest(BaseModel):
    grade: str = Field(..., max_length=20, json_schema_extra={"example": "三年级"})
    lesson: str = Field(..., max_length=100, json_schema_extra={"example": "富饶的西沙群岛"})
    requirements: str = Field(default="", max_length=500, json_schema_extra={"example": "重点修辞手法"})
    class_hours: str = Field(default="2", max_length=5)
    semester: str = Field(default="上", max_length=5)


class GenerateResponse(BaseModel):
    exam_analysis: str = Field(default="")
    peer_analysis: str = Field(default="")
    lesson_plan: str = Field(default="")
    teaching_guide: str = Field(default="")
    record_id: str = Field(default="")
    generated_blocks: list[dict] = Field(default_factory=list)
    teaching_evidence: list[dict] = Field(default_factory=list)
    missing_evidence: list[str] = Field(default_factory=list)
    citation_errors: list[str] = Field(default_factory=list)


class SectionGenerateRequest(GenerateRequest):
    record_id: str = Field(default="", max_length=120)
    lesson_plan: str = Field(default="", max_length=12000)
    exam_analysis: str = Field(default="", max_length=10000)
    peer_analysis: str = Field(default="", max_length=10000)


class TeachingEvidenceSearchRequest(BaseModel):
    grade: str = Field(default="", max_length=20)
    semester: str = Field(default="上", max_length=5)
    lesson: str = Field(default="", max_length=100)
    purpose: str = Field(default="lesson_plan", max_length=40)
    source_roles: list[str] = Field(default_factory=list)
    doc_types: list[str] = Field(default_factory=list)
    max_items: int = Field(default=20, ge=1, le=50)


class ReviseRequest(BaseModel):
    current_plan: str = Field(...)
    revision_request: str = Field(...)
    history: str = Field(default="")


class ReviseResponse(BaseModel):
    lesson_plan: str


class MindmapGenerateRequest(BaseModel):
    grade: str = ""
    lesson: str = Field(..., max_length=100)
    lesson_plan: str = Field(..., max_length=10000)
    teaching_guide: str = Field(default="", max_length=10000)
    analysis: str = Field(default="", max_length=10000)
    peer_reference: str = Field(default="", max_length=10000)


class MindmapGenerateResponse(BaseModel):
    lesson_mindmap_mermaid: str
    method_mindmap_mermaid: str
    lesson_outline: dict
    method_outline: dict


# ---- Public API ----

@app.get("/")
async def root():
    return RedirectResponse(url="/api/health")


# ---- 首次启动引导 ----

@app.get("/api/setup/status")
async def setup_status():
    from auth import load_users
    return {"needs_setup": len(load_users()) == 0}


@app.post("/api/setup/complete")
async def setup_complete(req: dict):
    from auth import load_users, register_user
    if len(load_users()) > 0:
        raise HTTPException(status_code=403, detail="系统已初始化")

    password = str(req.get("password", "")).strip()
    api_key = str(req.get("api_key", "")).strip().replace("\n", "").replace("\r", "")
    if len(password) < 4:
        raise HTTPException(status_code=400, detail="管理员密码至少4位")
    if not api_key:
        raise HTTPException(status_code=400, detail="请填写DeepSeek API Key")

    # 持久化 API Key：同时写 .env 和 data/api_key.json
    env_file = Path(__file__).resolve().parent.parent / ".env"
    # 保留现有 .env 中的其他配置
    existing = ""
    if env_file.exists():
        lines = env_file.read_text().split("\n")
        existing = "\n".join(l for l in lines if not l.startswith("DEEPSEEK_API_KEY=") and l.strip())
        if existing:
            existing += "\n"
    from security import atomic_write
    atomic_write(env_file, (existing + f"DEEPSEEK_API_KEY={api_key}\n").encode())
    import os as _os2
    _os2.chmod(env_file, 0o600)

    import json as _j
    key_file = Path(__file__).resolve().parent.parent / "data" / "api_key.json"
    atomic_write(key_file, _j.dumps({"api_key": api_key}).encode())
    _os2.chmod(key_file, 0o600)

    ok, msg = register_user("admin", password)
    if not ok:
        raise HTTPException(status_code=400, detail=msg)

    import rag
    rag.reload_keys(api_key)
    audit_log("admin", "admin", "setup_api_key", "", True, "初始化完成")
    return {"ok": True, "message": "初始化完成"}


@app.get("/api/setup/wifi-scan")
async def setup_wifi_scan():
    import subprocess
    try:
        r = subprocess.run(["nmcli","-t","-f","SSID,SIGNAL,SECURITY","dev","wifi","list","--rescan","yes"],
                         capture_output=True, text=True, timeout=15)
        nets = []
        for line in r.stdout.strip().split("\n"):
            p = line.split(":")
            if len(p) >= 2 and p[0].strip():
                nets.append({"ssid": p[0], "signal": int(p[1]) if len(p)>1 and p[1].isdigit() else 0,
                            "security": p[2] if len(p)>2 else ""})
        seen = set()
        uniq = [n for n in sorted(nets, key=lambda x: -x["signal"]) if not (n["ssid"] in seen or seen.add(n["ssid"]))]
        return {"networks": uniq[:20]}
    except Exception:
        return {"networks": [], "error": "WiFi扫描仅Linux盒子可用"}


@app.post("/api/setup/wifi-connect")
async def setup_wifi_connect(req: dict):
    ssid = str(req.get("ssid","")).strip()
    pw = str(req.get("password","")).strip()
    if not ssid:
        raise HTTPException(status_code=400, detail="请选择WiFi网络")
    import subprocess
    try:
        args = ["nmcli","dev","wifi","connect",ssid]
        if pw: args += ["password", pw]
        subprocess.run(args, capture_output=True, timeout=30, check=True)
        return {"ok": True, "ssid": ssid}
    except subprocess.CalledProcessError:
        raise HTTPException(status_code=400, detail="WiFi密码错误或信号弱")
    except Exception:
        raise HTTPException(status_code=400, detail="WiFi连接失败")

@app.get("/api/health")
async def health():
    return {"status": "ok", "version": VERSION}

@app.get("/api/health/deep")
async def health_deep(username: str = Depends(require_admin_or_reviewer)):
    """深度健康检查：验证API Key、Embedding模型、知识库"""
    try:
        from rag import call_deepseek
        call_deepseek("ping", "pong", temperature=0)
        api_ok = True
    except Exception:
        api_ok = False
    try:
        from search_engine import get_embedding_model
        get_embedding_model()
        model_ok = True
    except Exception:
        model_ok = False
    try:
        from search_engine import get_collection
        col = get_collection()
        chunk_count = col.count()
    except Exception:
        chunk_count = 0
    return {"status": "ok", "api_key_ok": api_ok, "model_ok": model_ok, "chunks": chunk_count}


@app.get("/api/textbooks")
async def textbooks():
    tree = []
    for grade in ["一年级","二年级","三年级","四年级","五年级","六年级"]:
        gd = GRADE_TEXTBOOKS.get(grade, {})
        sems = []
        for sn in ["上册","下册"]:
            sd = gd.get(sn, {})
            units = [{"name": un, "lessons": ls} for un, ls in sd.items()]
            sems.append({"name": sn, "units": units})
        tree.append({"grade": grade, "semesters": sems})
    return {"textbooks": tree}


# ---- Auth API ----

@app.post("/api/register")
async def register(req: AuthRequest, request: Request = None, username: str = Depends(require_auth)):
    """仅管理员可创建新用户（学校场景，关闭公开注册）"""
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可创建用户")
    client_ip = request.client.host if request else "127.0.0.1"
    if check_rate_limit(f"reg_{client_ip}", 5, 300):
        raise HTTPException(status_code=429, detail="注册过于频繁，请5分钟后再试")
    ok, msg = register_user(req.username, req.password)
    audit_log(username, role, "create_user", req.username, ok, msg)
    if not ok:
        raise HTTPException(status_code=400, detail=msg)
    return {"message": msg}


@app.post("/api/login")
async def login(req: AuthRequest, request: Request = None):
    client_ip = request.client.host if request else "127.0.0.1"
    if check_rate_limit(f"login_{client_ip}", 10, 60):
        raise HTTPException(status_code=429, detail="登录尝试过多，请1分钟后再试")
    token = login_user(req.username, req.password)
    if not token:
        raise HTTPException(status_code=401, detail="用户名或密码错误")
    role = get_user_role(req.username)
    return {"token": token, "username": req.username, "role": role}


@app.post("/api/logout")
async def logout(authorization: str = Header(default="")):
    if authorization.startswith("Bearer "):
        logout_user(authorization[len("Bearer "):])
    return {"message": "已退出"}


@app.get("/api/me")
async def me(authorization: str = Header(default="")):
    if authorization.startswith("Bearer "):
        username = get_user_from_token(authorization[len("Bearer "):])
        if username:
            role = get_user_role(username)
            return {"username": username, "role": role}
    return {"username": "", "role": ""}


# ---- Protected API ----

@app.post("/api/generate", response_model=GenerateResponse)
async def generate(
    req: GenerateRequest,
    request: Request,
    username: str = Depends(require_auth),
    x_accept_force_rag_error: str = Header(default="", alias="X-Accept-Force-Rag-Error")
):
    if not req.lesson.strip():
        raise HTTPException(status_code=400, detail="课题名称不能为空")

    # 生成端点速率限制
    if check_rate_limit(f"gen_{username}", 20, 60):
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    try:
        # 验收测试钩子：仅管理员可通过请求头触发
        if (
            LEKAI_ACCEPTANCE_MODE
            and x_accept_force_rag_error == "1"
            and get_user_role(username) == "admin"
        ):
            raise RuntimeError("验收用错误：模拟 RAG 调用失败")

        result = generate_lesson(
            grade=req.grade, lesson=req.lesson.strip(),
            requirements=req.requirements.strip(),
            class_hours=req.class_hours, semester=req.semester
        )
        record_id = save_history(username, req.grade, req.lesson.strip(), result)
        result["record_id"] = record_id
        return GenerateResponse(**result)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        import traceback, logging
        logging.getLogger("lekai").error("生成失败:\n%s", traceback.format_exc())
        raise HTTPException(status_code=500, detail="生成失败，请稍后重试")


def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {_json.dumps(data, ensure_ascii=False)}\n\n"


def _dedupe_evidence(items: list[dict]) -> list[dict]:
    seen: set[str] = set()
    out: list[dict] = []
    for item in items:
        eid = str(item.get("id") or "")
        if not eid or eid in seen:
            continue
        seen.add(eid)
        out.append(item)
    return out


def _append_history_metadata(username: str, record_id: str, fields: dict, bundle: dict) -> None:
    if not record_id:
        return
    detail = get_history_detail(username, record_id)
    if not detail:
        return
    evidence = _dedupe_evidence((detail.get("teaching_evidence") or []) + (bundle.get("evidence") or []))
    blocks = (detail.get("generated_blocks") or []) + (bundle.get("generated_blocks") or [])
    missing = list(dict.fromkeys((detail.get("missing_evidence") or []) + (bundle.get("missing") or [])))
    citation_errors = validate_block_citations(blocks, evidence_lookup(evidence))
    update = {
        **fields,
        "teaching_evidence": evidence,
        "generated_blocks": blocks,
        "missing_evidence": missing,
        "citation_errors": citation_errors,
    }
    update_history_fields(username, record_id, update)


def _appendix_markdown(evidence: list[dict]) -> str:
    normative = [e for e in evidence if e.get("source_role") == "normative"]
    method = [e for e in evidence if e.get("source_role") == "method_case"]

    def fmt(item: dict, prefix: str) -> str:
        label = DOC_TYPE_LABELS.get(item.get("doc_type", ""), item.get("doc_type", "材料"))
        page = f"第{item.get('page_num')}页" if item.get("page_num") else "页码待补"
        title = item.get("doc_title") or item.get("source_file_id") or "知识库材料"
        return f"- {prefix}：{label}《{title}》· {page}\n  {str(item.get('content') or '')[:180]}"

    lines = ["# 附录：依据与参考", "", "## 规范依据"]
    lines.extend(fmt(e, "依据") for e in normative[:20])
    if not normative:
        lines.append("- 暂无可追溯规范依据")
    lines.extend(["", "## 方法参考"])
    lines.extend(fmt(e, "参考") for e in method[:20])
    if not method:
        lines.append("- 暂无方法参考")
    return "\n".join(lines)


@app.post("/api/generate/stream")
async def generate_stream(
    req: GenerateRequest,
    username: str = Depends(require_auth),
    x_accept_force_rag_error: str = Header(default="", alias="X-Accept-Force-Rag-Error"),
):
    if not req.lesson.strip():
        raise HTTPException(status_code=400, detail="课题名称不能为空")
    if check_rate_limit(f"gen_{username}", 20, 60):
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    if (
        LEKAI_ACCEPTANCE_MODE
        and x_accept_force_rag_error == "1"
        and get_user_role(username) == "admin"
    ):
        raise HTTPException(status_code=400, detail="验收用错误：模拟 RAG 调用失败")

    def events():
        chunks: list[str] = []
        try:
            yield _sse("status", {"phase": "retrieving", "message": "正在检索知识库"})
            evidence_bundle = prepare_lesson_evidence(
                req.grade,
                req.lesson.strip(),
                req.semester,
                req.requirements.strip(),
            )
            yield _sse("status", {"phase": "generating", "message": "正在生成教案"})
            for chunk in stream_lesson_plan(
                grade=req.grade,
                lesson=req.lesson.strip(),
                requirements=req.requirements.strip(),
                class_hours=req.class_hours,
                semester=req.semester,
                evidence_bundle=evidence_bundle,
            ):
                chunks.append(chunk)
                yield _sse("token", {"text": chunk})

            lesson_plan = "".join(chunks).strip()
            if len(lesson_plan) < 500:
                raise RuntimeError("AI 返回教案过短，请稍后重试")
            blocks = build_generated_blocks(lesson_plan, evidence_bundle["normative"], evidence_bundle["method"])
            evidence = evidence_bundle["evidence"]
            result = {
                "exam_analysis": "",
                "peer_analysis": "",
                "lesson_plan": lesson_plan,
                "teaching_guide": "",
                "generated_blocks": blocks,
                "teaching_evidence": evidence,
                "missing_evidence": evidence_bundle["missing"],
                "citation_errors": validate_block_citations(blocks, evidence_lookup(evidence)),
            }
            record_id = save_history(username, req.grade, req.lesson.strip(), result)
            yield _sse("done", {"record_id": record_id, "missing_evidence": evidence_bundle["missing"]})
        except Exception as e:
            _log.exception("流式生成失败")
            yield _sse("error", {"detail": str(e) or "生成失败，请稍后重试"})

    return StreamingResponse(
        events(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/generate/exam")
async def generate_exam(req: SectionGenerateRequest, username: str = Depends(require_auth)):
    if not req.lesson.strip():
        raise HTTPException(status_code=400, detail="课题名称不能为空")
    try:
        bundle = generate_exam_analysis_bundle(req.grade, req.lesson.strip(), req.semester)
        text = bundle["text"]
        if req.record_id:
            _append_history_metadata(username, req.record_id, {"exam_analysis": text}, bundle)
        return {
            "exam_analysis": text,
            "teaching_evidence": bundle.get("evidence", []),
            "generated_blocks": bundle.get("generated_blocks", []),
            "missing_evidence": bundle.get("missing", []),
        }
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        _log.exception("考点分析生成失败")
        raise HTTPException(status_code=500, detail="考点分析生成失败，请稍后重试")


@app.post("/api/generate/peer")
async def generate_peer(req: SectionGenerateRequest, username: str = Depends(require_auth)):
    if not req.lesson.strip():
        raise HTTPException(status_code=400, detail="课题名称不能为空")
    try:
        bundle = generate_peer_analysis_bundle(req.grade, req.lesson.strip(), req.semester)
        text = bundle["text"]
        if req.record_id:
            _append_history_metadata(username, req.record_id, {"peer_analysis": text}, bundle)
        return {
            "peer_analysis": text,
            "teaching_evidence": bundle.get("evidence", []),
            "generated_blocks": bundle.get("generated_blocks", []),
            "missing_evidence": bundle.get("missing", []),
        }
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        _log.exception("同行参考生成失败")
        raise HTTPException(status_code=500, detail="同行参考生成失败，请稍后重试")


@app.post("/api/generate/guide")
async def generate_guide(req: SectionGenerateRequest, username: str = Depends(require_auth)):
    if not req.lesson_plan.strip():
        raise HTTPException(status_code=400, detail="请先生成教案")
    try:
        bundle = generate_teaching_guide_bundle(
            req.lesson.strip(),
            req.lesson_plan.strip(),
            req.exam_analysis,
            req.peer_analysis,
            grade=req.grade,
            semester=req.semester,
        )
        text = bundle["text"]
        if req.record_id:
            _append_history_metadata(username, req.record_id, {"teaching_guide": text}, bundle)
        return {
            "teaching_guide": text,
            "teaching_evidence": bundle.get("evidence", []),
            "generated_blocks": bundle.get("generated_blocks", []),
            "missing_evidence": bundle.get("missing", []),
        }
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        _log.exception("辅导说明生成失败")
        raise HTTPException(status_code=500, detail="辅导说明生成失败，请稍后重试")


@app.post("/api/revise", response_model=ReviseResponse)
async def revise(req: ReviseRequest, username: str = Depends(require_auth)):
    if not req.current_plan.strip():
        raise HTTPException(status_code=400, detail="请提供当前教案")
    try:
        new_plan = revise_lesson(req.current_plan, req.revision_request, req.history)
        return ReviseResponse(lesson_plan=new_plan)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        import traceback, logging
        logging.getLogger("lekai").error("修改失败:\n%s", traceback.format_exc())
        raise HTTPException(status_code=500, detail="修改失败，请稍后重试")


@app.get("/api/history")
async def history(username: str = Depends(require_auth)):
    """获取当前用户的生成历史"""
    return {"history": get_history(username)}


@app.get("/api/history/{record_id}")
async def history_detail(record_id: str, username: str = Depends(require_auth)):
    detail = get_history_detail(username, record_id)
    if not detail:
        raise HTTPException(status_code=404, detail="记录不存在")
    return detail


@app.post("/api/teaching-evidence/search")
async def api_teaching_evidence_search(req: TeachingEvidenceSearchRequest, username: str = Depends(require_auth)):
    evidence, missing = search_teaching_evidence(
        grade=req.grade,
        semester=req.semester,
        lesson=req.lesson,
        purpose=req.purpose,
        source_roles=req.source_roles or None,
        doc_types=req.doc_types or None,
        max_items=req.max_items,
    )
    return {"evidence": evidence, "missing": missing}


@app.get("/api/evidence/{evidence_id}")
async def api_evidence_detail(evidence_id: str, username: str = Depends(require_auth)):
    evidence = get_teaching_evidence_by_id(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="依据不存在")
    return evidence


@app.post("/api/history/{record_id}/mindmap")
async def history_save_mindmap(record_id: str, req: dict, username: str = Depends(require_auth)):
    """保存思维导图到历史记录"""
    lesson_mm = str(req.get("lesson_mindmap_mermaid", "")).strip()
    method_mm = str(req.get("method_mindmap_mermaid", "")).strip()

    if not lesson_mm and not method_mm:
        raise HTTPException(status_code=400, detail="思维导图内容不能为空")

    detail = get_history_detail(username, record_id)
    if not detail:
        raise HTTPException(status_code=404, detail="记录不存在")

    ok = save_history_mindmap(username, record_id, lesson_mm, method_mm)
    if not ok:
        raise HTTPException(status_code=500, detail="保存思维导图失败")

    return {"ok": True, "message": "思维导图已保存"}


# ---- Review API ----

class ReviewSubmitRequest(BaseModel):
    record_id: str

@app.post("/api/review/submit")
async def review_submit(req: ReviewSubmitRequest, username: str = Depends(require_auth)):
    ok = submit_for_review(username, req.record_id)
    if not ok:
        raise HTTPException(status_code=404, detail="记录不存在")
    return {"message": "已提交审核"}


# ---- Admin API ----

@app.get("/api/admin/dashboard")
async def admin_dashboard(username: str = Depends(require_admin_or_reviewer)):
    return get_dashboard_stats()

@app.get("/api/admin/reviews")
async def admin_reviews(username: str = Depends(require_admin_or_reviewer)):
    return {"reviews": get_review_queue()}

@app.get("/api/admin/reviews/{record_id}")
async def admin_review_detail(record_id: str, username: str = Depends(require_admin_or_reviewer)):
    detail = get_review_detail(record_id)
    if not detail:
        raise HTTPException(status_code=404, detail="审核记录不存在")
    return detail

@app.post("/api/admin/reviews/{record_id}/approve")
async def admin_approve(record_id: str, username: str = Depends(require_admin_or_reviewer)):
    ok = approve_review(record_id, username)
    if not ok:
        raise HTTPException(status_code=404, detail="审核记录不存在")
    return {"message": "已批准"}

@app.post("/api/admin/reviews/{record_id}/reject")
async def admin_reject(record_id: str, req: dict | None = None, username: str = Depends(require_admin_or_reviewer)):
    if req is None:
        req = {}
    comment = str(req.get("comment", ""))
    ok = reject_review(record_id, username, comment)
    if not ok:
        raise HTTPException(status_code=404, detail="审核记录不存在")
    return {"message": "已打回"}

@app.post("/api/admin/users/set-role")
async def admin_set_role(req: dict, username: str = Depends(require_admin)):
    target = req.get("username", "")
    role = req.get("role", "")
    ok = set_user_role(target, role)
    audit_log(username, "admin", "set_role", f"{target}->{role}", ok)
    if not ok:
        raise HTTPException(status_code=400, detail="设置失败")
    return {"message": f"已将 {target} 的角色设为 {role}"}


@app.post("/api/admin/users/import")
async def admin_import_users(req: dict, username: str = Depends(require_admin)):
    """CSV 批量导入教师账号。CSV 格式：username,password（用户名仅支持字母数字下划线连字符）"""
    import csv as _csv, io as _io
    csv_text = str(req.get("csv", "")).strip()
    if not csv_text:
        raise HTTPException(status_code=400, detail="请提供 CSV 内容")

    results = {"created": [], "failed": []}
    reader = _csv.reader(_io.StringIO(csv_text))
    for line_no, row in enumerate(reader, 1):
        if not row or (len(row) == 1 and not row[0].strip()):
            continue
        if row[0].strip().startswith("#"):
            continue
        if len(row) < 2:
            results["failed"].append({"line": line_no, "reason": "格式错误，需要 username,password"})
            continue
        u, p = row[0].strip(), row[1].strip()
        if not u or not p:
            results["failed"].append({"line": line_no, "reason": "用户名或密码为空"})
            continue
        ok, msg = register_user(u, p)
        audit_log(username, "admin", "import_user", u, ok, msg)
        if ok:
            results["created"].append(u)
        else:
            results["failed"].append({"line": line_no, "username": u, "reason": msg})

    return {"ok": True, "imported": len(results["created"]), "failed": len(results["failed"]), "results": results}


# ---- 单元规划 API ----

class UnitPlanRequest(BaseModel):
    grade: str = Field(default="三年级")
    unit: str = Field(default="第六单元")
    semester: str = Field(default="上")

@app.post("/api/unit-plan")
async def unit_plan(req: UnitPlanRequest, username: str = Depends(require_auth)):
    try:
        result = generate_unit_plan(req.grade, req.unit, req.semester)
        return {"unit_plan": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail="生成失败，请稍后重试")


# ---- 教研协作 API ----

@app.get("/api/collab/groups")
async def collab_groups(username: str = Depends(require_auth)):
    return {"groups": list_groups(), "my_groups": get_user_groups(username)}

@app.post("/api/collab/groups/create")
async def collab_create(req: dict, username: str = Depends(require_auth)):
    r = create_group(req.get("name", ""), username)
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r

@app.post("/api/collab/groups/join")
async def collab_join(req: dict, username: str = Depends(require_auth)):
    r = join_group(req.get("name", ""), username)
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r

@app.post("/api/collab/share")
async def collab_share(req: dict, username: str = Depends(require_auth)):
    r = share_plan(username, req.get("group", ""), req.get("grade", ""),
                   req.get("lesson", ""), req.get("record_id", ""))
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r

@app.get("/api/collab/groups/{group_name}/plans")
async def collab_plans(group_name: str, username: str = Depends(require_auth)):
    return {"plans": get_group_plans(group_name)}

@app.post("/api/collab/tasks/assign")
async def collab_assign(req: dict, username: str = Depends(require_auth)):
    r = assign_task(req.get("group", ""), req.get("assigned_to", ""),
                    req.get("lesson", ""), username)
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r

@app.get("/api/collab/groups/{group_name}/tasks")
async def collab_tasks(group_name: str, username: str = Depends(require_auth)):
    return {"tasks": get_group_tasks(group_name)}

@app.post("/api/collab/tasks/complete")
async def collab_complete(req: dict, username: str = Depends(require_auth)):
    r = complete_task(req.get("group", ""), req.get("task_index", 0))
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r

@app.post("/api/collab/comment")
async def collab_comment(req: dict, username: str = Depends(require_auth)):
    r = add_comment(req.get("group", ""), req.get("plan_index", 0),
                    username, req.get("text", ""))
    if not r["ok"]:
        raise HTTPException(status_code=400, detail=r["msg"])
    return r


# ---- 课后反思 API ----

class ReflectionRequest(BaseModel):
    lesson: str = Field(default="")
    lesson_plan: str = Field(default="")

@app.post("/api/reflect")
async def reflect(req: ReflectionRequest, username: str = Depends(require_auth)):
    try:
        result = generate_reflection(req.lesson, req.lesson_plan)
        return {"reflection": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail="生成失败，请稍后重试")


# ---- 健康检查 ----

@app.get("/api/admin/health")
async def admin_health(username: str = Depends(require_admin_or_reviewer)):
    return get_health()


# ---- 备份恢复 ----

@app.post("/api/admin/backup")
async def admin_backup(username: str = Depends(require_admin)):
    import time as _t
    buf = create_backup()
    ts = _t.strftime("%Y%m%d_%H%M%S")
    return StreamingResponse(buf, media_type="application/zip",
                            headers={"Content-Disposition": f"attachment; filename=lekai_backup_{ts}.zip"})


@app.post("/api/admin/restore")
async def admin_restore(file: UploadFile = FastAPIFile(...), username: str = Depends(require_admin)):
    data = await file.read()
    ok, msg = restore_backup(data)
    audit_log(username, "admin", "restore_backup", file.filename or "unknown", ok, msg)
    if not ok:
        raise HTTPException(status_code=400, detail=msg)
    return {"ok": True, "message": msg}


# ---- Prompt 在线配置 ----


@app.get("/api/admin/prompts")
async def admin_get_prompts(username: str = Depends(require_admin_or_reviewer)):
    if PROMPTS_FILE.exists():
        try:
            return _json.loads(PROMPTS_FILE.read_text())
        except (_json.JSONDecodeError, OSError):
            _log.warning("Prompt配置文件损坏，已重置: %s", PROMPTS_FILE.name)
    return {"chat_prompt": "", "audit_prompt": ""}


@app.post("/api/admin/prompts")
async def admin_set_prompts(req: dict, username: str = Depends(require_admin)):
    cur = {}
    if PROMPTS_FILE.exists():
        try:
            cur = _json.loads(PROMPTS_FILE.read_text())
        except (_json.JSONDecodeError, OSError):
            _log.warning("Prompt配置文件损坏，备份后重置: %s", PROMPTS_FILE.name)
            try:
                import shutil
                shutil.copy2(PROMPTS_FILE, PROMPTS_FILE.with_suffix(".corrupted"))
            except Exception:
                pass
    for key in ("chat_prompt", "audit_prompt"):
        if key in req:
            cur[key] = str(req[key])[:5000]
    from security import atomic_write
    atomic_write(PROMPTS_FILE, _json.dumps(cur, ensure_ascii=False, indent=2).encode())
    audit_log(username, "admin", "update_prompt", ", ".join(req.keys()))
    return {"ok": True, "message": "提示词已更新，立即生效"}


# ---- 教案反馈 ----

class FeedbackRequest(BaseModel):
    plan_id: str
    grade: str = ""
    lesson: str = ""
    rating: int = 0
    useful_refs: list[str] = []
    tags: list[str] = []
    comment: str = ""

@app.post("/api/feedback")
async def api_feedback(req: FeedbackRequest, username: str = Depends(require_auth)):
    return submit_feedback(username, req.plan_id, req.grade, req.lesson,
                           req.rating, req.useful_refs, req.tags, req.comment)

@app.get("/api/feedback/{plan_id}")
async def api_get_feedback(plan_id: str, username: str = Depends(require_auth)):
    fb = get_feedback(plan_id)
    if not fb:
        raise HTTPException(status_code=404, detail="无反馈记录")
    return fb

@app.get("/api/admin/feedback-stats")
async def admin_feedback_stats(username: str = Depends(require_admin_or_reviewer)):
    return get_feedback_stats()


# ---- 段落标注 ----

class AnnotationRequest(BaseModel):
    review_id: str
    section: str
    type: str  # praise | improve | note
    text: str

@app.post("/api/admin/annotations")
async def api_add_annotation(req: AnnotationRequest, username: str = Depends(require_admin_or_reviewer)):
    return add_annotation(req.review_id, username, req.section, req.type, req.text)

@app.get("/api/admin/annotations/{review_id}")
async def api_get_annotations(review_id: str, username: str = Depends(require_admin_or_reviewer)):
    return {"annotations": get_annotations(review_id)}

@app.get("/api/admin/annotation-stats")
async def admin_annotation_stats(username: str = Depends(require_admin_or_reviewer)):
    return get_annotation_stats()


# ---- 思维导图生成 ----

@app.post("/api/mindmap/generate", response_model=MindmapGenerateResponse)
async def generate_mindmap(
    req: MindmapGenerateRequest,
    username: str = Depends(require_auth),
    x_accept_force_mindmap_error: str = Header(default="", alias="X-Accept-Force-Mindmap-Error")
):
    if not req.lesson.strip():
        raise HTTPException(status_code=400, detail="课题名称不能为空")
    if not req.lesson_plan.strip():
        raise HTTPException(status_code=400, detail="教案内容不能为空")
    if len(req.lesson_plan.strip()) < 100:
        raise HTTPException(status_code=400, detail="教案内容过短，无法生成思维导图")

    # 速率限制（mindmap 生成成本较高，频率限制低于教案生成）
    if check_rate_limit(f"mindmap_{username}", 10, 60):
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")

    try:
        # 验收测试钩子：仅管理员可通过请求头触发
        if (
            LEKAI_ACCEPTANCE_MODE
            and x_accept_force_mindmap_error == "1"
            and get_user_role(username) == "admin"
        ):
            raise RuntimeError("验收用错误：模拟思维导图生成失败")

        result = generate_dual_mindmap(req)
        return MindmapGenerateResponse(**result)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        import traceback, logging
        logging.getLogger("lekai").error("思维导图生成失败:\n%s", traceback.format_exc())
        raise HTTPException(status_code=500, detail="思维导图生成失败，请稍后重试")


# ---- 教案评价（老师上传自己的教案，AI对照知识库评价） ----

@app.post("/api/evaluate")
async def evaluate_plan(file: UploadFile = FastAPIFile(...), username: str = Depends(require_auth)):
    """上传教案文件，AI对照知识库中的优秀教案进行评价"""
    fn = file.filename or "untitled"
    ext = Path(fn).suffix.lower()
    if ext not in (".md", ".txt", ".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .md / .txt / .docx 格式")
    # 上传安全：大小限制 + 安全文件名
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件不能超过5MB")

    # 解析文档内容（复用已读取的 data 变量）
    if ext == ".docx":
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(data))
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        content = data.decode("utf-8", errors="ignore")

    if len(content) < 200:
        raise HTTPException(status_code=400, detail="教案内容过短，至少200字符")

    # 提取课题名
    import re
    match = re.search(r'《(.+?)》', content)
    lesson_name = match.group(1) if match else Path(fn).stem

    # 检索知识库中的同类优秀教案
    from rag import retrieve_structured, call_deepseek
    context, _ = retrieve_structured(f"{lesson_name} 教案", top_k=8)

    # 调用评价 Prompt
    from prompts import EVALUATE_SYSTEM, EVALUATE_USER
    eval_prompt = EVALUATE_USER.format(uploaded_plan=content[:6000], context=context, lesson_name=lesson_name)
    try:
        evaluation = call_deepseek(EVALUATE_SYSTEM, eval_prompt, temperature=0.3)
    except Exception:
        raise HTTPException(status_code=500, detail="评价生成失败，请稍后重试")

    return {"evaluation": evaluation, "lesson_name": lesson_name}


# ---- 教案导出 ----

@app.get("/api/export/{plan_id}")
async def export_plan(plan_id: str, format: str = "md", include_mindmap: str = "false",
                      username: str = Depends(require_auth)):
    """导出教案为 md / docx"""
    from auth import get_history_detail
    detail = get_history_detail(username, plan_id)
    if not detail:
        raise HTTPException(status_code=404, detail="教案不存在")

    plan_text = detail.get("lesson_plan", "")
    guide_text = detail.get("teaching_guide", "")
    lesson = detail.get("lesson", "教案")
    grade = detail.get("grade", "")
    include_mm = include_mindmap.lower() == "true"

    from urllib.parse import quote

    if format == "md":
        full = plan_text
        if guide_text:
            full += "\n\n---\n\n" + guide_text
        full += "\n\n---\n\n" + _appendix_markdown(detail.get("teaching_evidence") or [])

        if include_mm:
            lesson_mm = detail.get("lesson_mindmap_mermaid", "")
            method_mm = detail.get("method_mindmap_mermaid", "")
            if lesson_mm or method_mm:
                full += "\n\n---\n\n# 附录：教案思维导图\n\n```mermaid\n" + (lesson_mm or "> 暂无思维导图") + "\n```"
                full += "\n\n# 附录：备课方法思维导图\n\n```mermaid\n" + (method_mm or "> 暂无思维导图") + "\n```"

        from fastapi.responses import Response
        safe_fn = quote(f"{lesson}_教案.md")
        return Response(content=full.encode("utf-8"), media_type="text/markdown",
                       headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_fn}"})

    elif format == "docx":
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading(f"《{lesson}》教案", 0)
        if grade:
            doc.add_paragraph(f"年级：{grade}", style="Subtitle")

        for line in plan_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], 1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], 3)
            elif line.strip():
                doc.add_paragraph(line.strip())

        if guide_text:
            doc.add_page_break()
            doc.add_heading("教案辅导说明", 1)
            for line in guide_text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], 1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 2)
                elif line.strip():
                    doc.add_paragraph(line.strip())

        doc.add_page_break()
        doc.add_heading("附录：依据与参考", 1)
        normative, method = split_evidence(detail.get("teaching_evidence") or [])
        doc.add_heading("规范依据", 2)
        if normative:
            for item in normative[:20]:
                label = DOC_TYPE_LABELS.get(item.get("doc_type", ""), item.get("doc_type", "材料"))
                page = f"第{item.get('page_num')}页" if item.get("page_num") else "页码待补"
                doc.add_paragraph(f"依据：{label}《{item.get('doc_title') or '知识库材料'}》 · {page}")
        else:
            doc.add_paragraph("暂无可追溯规范依据")
        doc.add_heading("方法参考", 2)
        if method:
            for item in method[:20]:
                label = DOC_TYPE_LABELS.get(item.get("doc_type", ""), item.get("doc_type", "材料"))
                page = f"第{item.get('page_num')}页" if item.get("page_num") else "页码待补"
                doc.add_paragraph(f"参考：{label}《{item.get('doc_title') or '知识库材料'}》 · {page}")
        else:
            doc.add_paragraph("暂无方法参考")

        if include_mm:
            lesson_mm = detail.get("lesson_mindmap_mermaid", "")
            method_mm = detail.get("method_mindmap_mermaid", "")
            if lesson_mm or method_mm:
                doc.add_page_break()
                doc.add_heading("附录：教案思维导图（Mermaid 源码）", 1)
                doc.add_paragraph(lesson_mm or "暂无思维导图")
                doc.add_heading("附录：备课方法思维导图（Mermaid 源码）", 1)
                doc.add_paragraph(method_mm or "暂无思维导图")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_fn = quote(f"{lesson}_教案.docx")
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_fn}"})

    raise HTTPException(status_code=400, detail="仅支持 md / docx 格式")


# ---- 教案入库（管理员上传教案文档） ----

def _safe_filename(value: str, fallback: str = "material") -> str:
    import re
    cleaned = re.sub(r'[《》\s/:*?"<>|\\]+', "", value or "").strip(".")
    return (cleaned or fallback)[:100]


def _grade_dir_name(grade_name: str) -> str:
    grade_dirs = {
        "一年级": "grade-1", "二年级": "grade-2", "三年级": "grade-3",
        "四年级": "grade-4", "五年级": "grade-5", "六年级": "grade-6",
    }
    return grade_dirs.get(grade_name, "")


def _infer_grade_from_text(text: str) -> str:
    import re
    m = re.search(r'(一年级|二年级|三年级|四年级|五年级|六年级)', text)
    if m:
        return m.group(1)
    m = re.search(r'([1-6一二三四五六])\s*年级', text)
    if not m:
        return ""
    mp = {"1": "一年级", "一": "一年级", "2": "二年级", "二": "二年级",
          "3": "三年级", "三": "三年级", "4": "四年级", "四": "四年级",
          "5": "五年级", "五": "五年级", "6": "六年级", "六": "六年级"}
    return mp.get(m.group(1), "")


def _infer_semester_from_text(text: str) -> str:
    return "下" if "下册" in text or "下学期" in text else "上"


def _run_ingest_and_refresh(timeout: int = 240) -> tuple[bool, str]:
    import subprocess, sys
    result = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent.parent / "scripts" / "ingest_knowledge.py")],
        capture_output=True, text=True, timeout=timeout)
    if result.returncode != 0:
        return False, (result.stderr or result.stdout or "")[:1000]
    try:
        from search_engine import refresh_index
        refresh_index()
    except Exception:
        _log.exception("入库后BM25索引刷新失败")
    return True, result.stdout[-1000:]


def _ocr_tsv_to_text_and_bbox(tsv: str, page_width: int, page_height: int) -> tuple[str, list[int], float]:
    import csv
    from io import StringIO
    words: list[dict] = []
    confidences: list[float] = []
    for row in csv.DictReader(StringIO(tsv), delimiter="\t"):
        text = (row.get("text") or "").strip()
        if not text:
            continue
        try:
            conf = float(row.get("conf") or -1)
            left = int(float(row.get("left") or 0))
            top = int(float(row.get("top") or 0))
            width = int(float(row.get("width") or 0))
            height = int(float(row.get("height") or 0))
            block = int(row.get("block_num") or 0)
            par = int(row.get("par_num") or 0)
            line = int(row.get("line_num") or 0)
        except ValueError:
            continue
        if conf >= 0:
            confidences.append(conf)
        words.append({
            "text": text, "left": left, "top": top, "right": left + width, "bottom": top + height,
            "block": block, "par": par, "line": line,
        })
    if not words:
        return "", [0, 0, page_width, page_height], 0.0
    lines: dict[tuple[int, int, int], list[dict]] = {}
    for word in words:
        lines.setdefault((word["block"], word["par"], word["line"]), []).append(word)
    ordered_lines = []
    for key, line_words in sorted(lines.items(), key=lambda item: (min(w["top"] for w in item[1]), min(w["left"] for w in item[1]))):
        ordered = sorted(line_words, key=lambda w: w["left"])
        ordered_lines.append("".join(w["text"] for w in ordered))
    left = max(0, min(w["left"] for w in words))
    top = max(0, min(w["top"] for w in words))
    right = min(page_width, max(w["right"] for w in words))
    bottom = min(page_height, max(w["bottom"] for w in words))
    avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
    return "\n".join(ordered_lines), [left, top, right, bottom], avg_conf


def _ocr_pdf_to_markdown_files(
    *,
    pdf_bytes: bytes,
    filename: str,
    doc_type: str,
    source_role: str,
) -> tuple[str, str, int, int]:
    """Render scanned PDF pages, OCR them, and write page-level markdown files."""
    import hashlib
    import shutil
    import subprocess
    import tempfile

    if not shutil.which("tesseract"):
        raise HTTPException(status_code=400, detail="本机未安装 tesseract，无法处理扫描 PDF")
    try:
        import pypdfium2 as pdfium
    except Exception as exc:
        raise HTTPException(status_code=400, detail="缺少 pypdfium2，无法渲染 PDF 页图") from exc

    source_hash = hashlib.sha256(pdf_bytes).hexdigest()
    stem = Path(filename).stem
    grade_name = _infer_grade_from_text(filename) or "一年级"
    semester = _infer_semester_from_text(filename)
    lesson_name = stem
    safe_stem = _safe_filename(stem, "pdf")
    project_root = Path(__file__).resolve().parent.parent
    uploads_dir = project_root / "data" / "uploads" / "pdf"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = uploads_dir / f"{source_hash[:16]}_{safe_stem}.pdf"
    pdf_path.write_bytes(pdf_bytes)

    page_dir = project_root / "data" / "source_pages" / source_hash[:16]
    page_dir.mkdir(parents=True, exist_ok=True)

    grade_dir = _grade_dir_name(grade_name)
    kb_root = project_root / "knowledge-base"
    dest_dir = (kb_root / grade_dir if grade_dir else kb_root) / "ocr" / safe_stem
    dest_dir.mkdir(parents=True, exist_ok=True)
    for old in dest_dir.glob("*.md"):
        old.unlink()

    tmp_root = project_root / "tmp_pdf_check"
    tmp_root.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(dir=str(tmp_root)):
        pdf = pdfium.PdfDocument(str(pdf_path))
        total_pages = len(pdf)
        page_count = min(total_pages, PDF_OCR_MAX_PAGES)
        written = 0
        for page_index in range(page_count):
            page = pdf[page_index]
            bitmap = page.render(scale=PDF_OCR_DPI_SCALE)
            image = bitmap.to_pil()
            image_path = page_dir / f"page_{page_index + 1:04d}.png"
            image.save(image_path)
            tsv = subprocess.run(
                ["tesseract", str(image_path), "stdout", "-l", PDF_OCR_LANG, "--psm", "6", "tsv"],
                capture_output=True, text=True, timeout=90,
            )
            if tsv.returncode != 0:
                _log.warning("OCR 第 %s 页失败: %s", page_index + 1, (tsv.stderr or "")[:300])
                continue
            text, bbox, confidence = _ocr_tsv_to_text_and_bbox(tsv.stdout, image.width, image.height)
            if len(text.strip()) < 20:
                continue
            page_url = f"/api/source-pages/{source_hash[:16]}/page_{page_index + 1:04d}.png"
            md = (
                "---\n"
                f"title: \"{stem} 第{page_index + 1}页 OCR\"\n"
                f"doc_title: \"{stem}\"\n"
                f"grade: \"{grade_name}\"\n"
                f"semester: \"{semester}\"\n"
                "unit: \"\"\n"
                f"lesson: \"{lesson_name}\"\n"
                "type: \"扫描PDF OCR\"\n"
                f"doc_type: \"{doc_type}\"\n"
                f"source_role: \"{source_role}\"\n"
                f"source_file_id: \"{source_hash[:16]}\"\n"
                f"source_file_name: \"{filename}\"\n"
                f"source_hash: \"{source_hash}\"\n"
                f"page_num: {page_index + 1}\n"
                f"page_width: {image.width}\n"
                f"page_height: {image.height}\n"
                f"bbox: {bbox}\n"
                f"page_image_url: \"{page_url}\"\n"
                f"ocr_confidence: {confidence:.2f}\n"
                "tags: [\"扫描PDF\", \"OCR\"]\n"
                "---\n\n"
                f"# 《{lesson_name}》第{page_index + 1}页 OCR\n\n"
                "## 教学过程\n"
                f"{text.strip()}\n"
            )
            (dest_dir / f"page_{page_index + 1:04d}.md").write_text(md, encoding="utf-8")
            written += 1
        return lesson_name, grade_name, total_pages, written

@app.post("/api/admin/upload-lesson")
async def admin_upload_lesson(
    file: UploadFile = FastAPIFile(...),
    doc_type: str = Form(default=""),
    source_role: str = Form(default=""),
    username: str = Depends(require_admin_or_reviewer)
):
    """上传教学材料，支持 .md/.txt/.docx 和扫描 PDF OCR 入库。"""
    import re
    doc_type = doc_type.strip()
    source_role = source_role.strip()
    if doc_type not in DOC_TYPE_LABELS:
        raise HTTPException(status_code=400, detail="请选择有效的文档类型")
    if source_role not in ("normative", "method_case"):
        raise HTTPException(status_code=400, detail="请选择材料角色")
    expected_role = "normative" if doc_type in NORMATIVE_DOC_TYPES else "method_case"
    if source_role != expected_role:
        raise HTTPException(status_code=400, detail=f"{DOC_TYPE_LABELS[doc_type]} 应归为 {'规范依据' if expected_role == 'normative' else '方法参考'}")

    fn = file.filename or "untitled"
    ext = Path(fn).suffix.lower()
    if ext not in (".md", ".txt", ".docx", ".pdf"):
        raise HTTPException(status_code=400, detail="仅支持 .md / .txt / .docx / .pdf 格式")
    # 上传安全：大小限制 + 安全文件名
    data = await file.read()
    max_size = PDF_UPLOAD_MAX_MB * 1024 * 1024 if ext == ".pdf" else 5 * 1024 * 1024
    if len(data) > max_size:
        raise HTTPException(status_code=400, detail=f"文件不能超过{max_size // 1024 // 1024}MB")

    if ext == ".pdf":
        try:
            lesson_name, grade_name, total_pages, ocr_pages = _ocr_pdf_to_markdown_files(
                pdf_bytes=data,
                filename=fn,
                doc_type=doc_type,
                source_role=source_role,
            )
        except HTTPException:
            raise
        except Exception as exc:
            _log.exception("PDF OCR 入库失败")
            raise HTTPException(status_code=400, detail=f"PDF OCR 失败: {exc}") from exc

        ok, detail = _run_ingest_and_refresh(timeout=300)
        if not ok:
            _log.error("PDF OCR 后入库失败: %s", detail[:500])
            return {"ok": False, "lesson": lesson_name, "message": f"《{lesson_name}》OCR 完成但入库失败，请查看服务端日志"}
        role = get_user_role(username)
        audit_log(username, role, "ingest_pdf_ocr", lesson_name, True, f"{ocr_pages}/{total_pages} pages")
        cap = f"（已处理前 {ocr_pages}/{total_pages} 页）" if ocr_pages < total_pages else f"（已处理 {ocr_pages} 页）"
        return {"ok": True, "lesson": lesson_name, "message": f"《{lesson_name}》PDF OCR 已入库{cap}"}

    # 解析文档内容（复用已读取的 data 变量）
    if ext == ".docx":
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(data))
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        content = data.decode("utf-8", errors="ignore")

    if len(content) < 100:
        raise HTTPException(status_code=400, detail="文档内容过短（至少100字符）")

    # 提取课题名做文件名（在 DeepSeek 格式化之前，避免测试钩子浪费 API 调用）
    match = re.search(r'《(.+?)》', content)
    lesson_name = match.group(1) if match else Path(fn).stem

    # 测试钩子：文件名含 force_ingest_fail 时模拟入库失败（仅验收模式下生效）
    if LEKAI_ACCEPTANCE_MODE and "force_ingest_fail" in fn.lower() and get_user_role(username) in ("admin", "reviewer"):
        return {"ok": False, "lesson": lesson_name, "message": f"验收用：模拟入库失败"}

    def _has_frontmatter(text: str) -> bool:
        return bool(re.match(r'^\s*---\s*\n[\s\S]*?\n---\s*\n', text))

    def _extract_grade(text: str) -> str:
        m = re.search(r'(一年级|二年级|三年级|四年级|五年级|六年级)', text)
        if m:
            return m.group(1)
        m = re.search(r'([1-6一二三四五六])\s*年级', text)
        if not m:
            return ""
        mp = {"1": "一年级", "一": "一年级", "2": "二年级", "二": "二年级",
              "3": "三年级", "三": "三年级", "4": "四年级", "四": "四年级",
              "5": "五年级", "五": "五年级", "6": "六年级", "六": "六年级"}
        return mp.get(m.group(1), "")

    def _extract_meta_value(text: str, key: str) -> str:
        m = re.search(rf'^\s*{re.escape(key)}\s*:\s*(.+?)\s*$', text, re.MULTILINE)
        return m.group(1).strip().strip('"').strip("'") if m else ""

    def _upsert_frontmatter(text: str, fields: dict[str, str]) -> str:
        match = re.match(r'^(\s*---\s*\n)([\s\S]*?)(\n---\s*\n)', text)
        if not match:
            lines = ["---"] + [f"{k}: {v}" for k, v in fields.items()] + ["---", "", text.strip(), ""]
            return "\n".join(lines)
        body = match.group(2)
        for key, value in fields.items():
            pattern = rf'^\s*{re.escape(key)}\s*:.*$'
            replacement = f"{key}: {value}"
            if re.search(pattern, body, re.MULTILINE):
                body = re.sub(pattern, replacement, body, flags=re.MULTILINE)
            else:
                body = body.rstrip() + f"\n{replacement}"
        return text[:match.start(2)] + body + text[match.end(2):]

    def _ensure_frontmatter(text: str, lesson: str) -> tuple[str, str, str]:
        grade = _extract_meta_value(text, "grade") or _extract_grade(text)
        sem = _extract_meta_value(text, "semester")
        if not sem:
            sem = "下" if "下册" in text or "下学期" in text else "上"
        lesson_meta = _extract_meta_value(text, "lesson") or lesson
        lesson_type = _extract_meta_value(text, "type") or ("规范材料" if source_role == "normative" else "阅读课")
        if _has_frontmatter(text):
            patched = _upsert_frontmatter(text, {"doc_type": doc_type, "source_role": source_role})
            return patched, grade, lesson_meta
        fm = (
            "---\n"
            f"grade: {grade}\n"
            f"semester: {sem}\n"
            "unit:\n"
            f"lesson: {lesson_meta}\n"
            f"type: {lesson_type}\n"
            f"doc_type: {doc_type}\n"
            f"source_role: {source_role}\n"
            "class_hours: 2\n"
            "tags: []\n"
            "source: 上传入库\n"
            "curated_by: LeKai\n"
            "---\n\n"
        )
        return fm + text.strip() + "\n", grade, lesson_meta

    # 标准 Markdown 直接入库；规范材料也直接保留原文，不能用“教案格式化”改写原始依据。
    if _has_frontmatter(content):
        formatted = content
    elif source_role == "normative":
        formatted = content
    else:
        try:
            from rag import call_deepseek
            fmt_sys = "你是教案格式化助手。将原始教案内容整理为规范Markdown格式。保留全部教学要点。"
            fmt_user = (
                "请将以下原始教案整理为带 YAML frontmatter 的标准 Markdown 教案。"
                "frontmatter 必须包含 grade、semester、unit、lesson、type、class_hours、tags。\n\n"
                f"{content[:8000]}"
            )
            formatted = call_deepseek(fmt_sys, fmt_user)
        except Exception:
            formatted = content  # API不可用时直接用原文

    # 用格式化后的内容重新提取课题名
    match = re.search(r'《(.+?)》', formatted)
    lesson_name = match.group(1) if match else lesson_name
    formatted, grade_name, lesson_name = _ensure_frontmatter(formatted, lesson_name)

    # 保存到 knowledge-base/
    import re as _re
    safe_name = _re.sub(r'[《》\s/:*?"<>|]', '', lesson_name)[:100]
    grade_dirs = {
        "一年级": "grade-1", "二年级": "grade-2", "三年级": "grade-3",
        "四年级": "grade-4", "五年级": "grade-5", "六年级": "grade-6",
    }
    kb_root = Path(__file__).resolve().parent.parent / "knowledge-base"
    dest_dir = kb_root / grade_dirs[grade_name] if grade_name in grade_dirs else kb_root
    dest = dest_dir / f"{safe_name}.md"
    from security import atomic_write
    atomic_write(dest, formatted.encode("utf-8"))

    # 触发入库
    import subprocess, sys
    try:
        result = subprocess.run(
            [sys.executable, str(Path(__file__).resolve().parent.parent / "scripts" / "ingest_knowledge.py")],
            capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            import logging
            logging.getLogger("lekai").error("入库失败: %s", (result.stderr or result.stdout or "")[:500])
            return {"ok": False, "lesson": lesson_name, "message": f"《{lesson_name}》入库失败，请查看服务端日志"}
        # 刷新主进程 BM25 索引
        try:
            from search_engine import refresh_index
            refresh_index()
        except Exception:
            _log.exception("入库后BM25索引刷新失败")
        role = get_user_role(username)
        audit_log(username, role, "ingest_lesson", lesson_name)
        return {"ok": True, "lesson": lesson_name, "message": f"《{lesson_name}》已入库"}
    except Exception:
        return {"ok": False, "lesson": lesson_name, "message": f"《{lesson_name}》已保存，请手动运行入库脚本"}


# ---- 设备信息 ----

@app.get("/api/admin/device-info")
async def admin_device_info(username: str = Depends(require_admin_or_reviewer)):
    import shutil
    from config import get_device_mac
    proj = Path(__file__).resolve().parent.parent
    mac = get_device_mac()
    disk = shutil.disk_usage(proj)
    lic_file = proj / ".license"
    lic_status = "已授权" if lic_file.exists() else "未授权"

    return {
        "mac": mac,
        "disk_total_gb": round(disk.total / 1024**3, 1),
        "disk_free_gb": round(disk.free / 1024**3, 1),
        "disk_used_pct": round((disk.used / disk.total) * 100, 1),
        "license": lic_status,
        "version": VERSION,
    }


@app.get("/api/admin/evidence-coverage")
async def admin_evidence_coverage(
    grade: str = "",
    semester: str = "",
    lesson: str = "",
    username: str = Depends(require_admin_or_reviewer),
):
    """Check whether lessons have enough normative evidence and method references."""
    from search_engine import get_collection

    def summarize(items: list[dict]) -> dict:
        normative = [e for e in items if e.get("source_role") == "normative"]
        method = [e for e in items if e.get("source_role") == "method_case"]
        doc_counts: dict[str, int] = {}
        for item in items:
            doc_type = str(item.get("doc_type") or "unknown")
            doc_counts[doc_type] = doc_counts.get(doc_type, 0) + 1
        has_textbook = any(e.get("doc_type") == "textbook" for e in normative)
        has_standard_or_exam = any(e.get("doc_type") in {"curriculum_standard", "exam_outline", "unit_goal", "exam_material"} for e in normative)
        has_method = bool(method)
        missing: list[str] = []
        if not has_textbook:
            missing.append("缺少教材")
        if not has_standard_or_exam:
            missing.append("缺少课标/考纲/单元目标/考试资料")
        if not has_method:
            missing.append("缺少方法案例")
        return {
            "normative_count": len(normative),
            "method_count": len(method),
            "doc_type_counts": doc_counts,
            "has_textbook": has_textbook,
            "has_standard_or_exam": has_standard_or_exam,
            "has_method_case": has_method,
            "missing": missing,
            "status": "ready" if not missing else "insufficient",
        }

    if lesson.strip():
        evidence, missing = search_teaching_evidence(
            grade=grade,
            semester=semester or "上",
            lesson=lesson.strip(),
            purpose="guidance",
            max_items=50,
        )
        summary = summarize(evidence)
        summary["missing"] = list(dict.fromkeys(summary["missing"] + missing))
        return {"grade": grade, "semester": semester, "lesson": lesson.strip(), **summary, "evidence": evidence}

    col = get_collection()
    if col.count() == 0:
        return {"lessons": [], "total_lessons": 0}
    result = col.get(include=["metadatas"], limit=2000)
    grouped: dict[str, list[dict]] = {}
    for meta in result.get("metadatas") or []:
        if grade and meta.get("grade") != grade:
            continue
        if semester and meta.get("semester") not in ("", semester):
            continue
        lesson_name = str(meta.get("lesson") or meta.get("doc_title") or "未标注课题")
        grouped.setdefault(lesson_name, []).append(meta)
    lessons = [
        {"lesson": lesson_name, **summarize(items)}
        for lesson_name, items in sorted(grouped.items())
    ]
    return {"grade": grade, "semester": semester, "lessons": lessons, "total_lessons": len(lessons)}


@app.get("/api/admin/evidence-gaps")
async def admin_evidence_gaps(
    grade: str = "",
    lesson: str = "",
    username: str = Depends(require_admin_or_reviewer),
):
    """Summarize evidence gaps found in real generated lesson histories."""
    from auth import HISTORY_DIR

    total_records = 0
    records_with_gaps = 0
    citation_error_records = 0
    missing_counts: dict[str, int] = {}
    insufficient_block_counts: dict[str, int] = {}
    lesson_map: dict[str, dict] = {}
    recent_records: list[dict] = []

    def add_count(target: dict[str, int], key: str) -> None:
        if not key:
            return
        target[key] = target.get(key, 0) + 1

    if not HISTORY_DIR.exists():
        return {
            "total_records": 0,
            "records_with_gaps": 0,
            "citation_error_records": 0,
            "missing_counts": {},
            "insufficient_block_counts": {},
            "lessons": [],
            "recent_records": [],
        }

    for user_dir in HISTORY_DIR.iterdir():
        if not user_dir.is_dir():
            continue
        for path in sorted(user_dir.glob("*.json"), reverse=True):
            try:
                data = _json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                continue
            record_grade = str(data.get("grade") or "")
            record_lesson = str(data.get("lesson") or "")
            if grade and record_grade != grade:
                continue
            if lesson and lesson not in record_lesson:
                continue

            total_records += 1
            missing = [str(item) for item in (data.get("missing_evidence") or []) if str(item).strip()]
            blocks = data.get("generated_blocks") or []
            insufficient_blocks = [
                str(block.get("block_type") or block.get("id") or "unknown")
                for block in blocks
                if isinstance(block, dict) and block.get("evidence_status") == "insufficient"
            ]
            citation_errors = [str(item) for item in (data.get("citation_errors") or []) if str(item).strip()]
            has_gap = bool(missing or insufficient_blocks or citation_errors)
            if has_gap:
                records_with_gaps += 1
            if citation_errors:
                citation_error_records += 1
            for item in missing:
                add_count(missing_counts, item)
            for block_type in insufficient_blocks:
                add_count(insufficient_block_counts, block_type)

            key = f"{record_grade}\t{record_lesson}"
            entry = lesson_map.setdefault(key, {
                "grade": record_grade,
                "lesson": record_lesson,
                "total_records": 0,
                "records_with_gaps": 0,
                "citation_error_records": 0,
                "missing_counts": {},
                "insufficient_block_counts": {},
                "last_seen": "",
            })
            entry["total_records"] += 1
            if has_gap:
                entry["records_with_gaps"] += 1
            if citation_errors:
                entry["citation_error_records"] += 1
            for item in missing:
                add_count(entry["missing_counts"], item)
            for block_type in insufficient_blocks:
                add_count(entry["insufficient_block_counts"], block_type)
            entry["last_seen"] = max(str(entry.get("last_seen") or ""), str(data.get("timestamp") or ""))

            if has_gap:
                recent_records.append({
                    "id": path.stem,
                    "username": user_dir.name,
                    "grade": record_grade,
                    "lesson": record_lesson,
                    "timestamp": data.get("timestamp") or "",
                    "missing_evidence": missing,
                    "insufficient_blocks": insufficient_blocks,
                    "citation_errors": citation_errors,
                })

    lessons = sorted(
        lesson_map.values(),
        key=lambda item: (item["records_with_gaps"], item["total_records"], item["last_seen"]),
        reverse=True,
    )
    recent_records = sorted(recent_records, key=lambda item: item.get("timestamp") or "", reverse=True)[:20]
    return {
        "grade": grade,
        "lesson": lesson,
        "total_records": total_records,
        "records_with_gaps": records_with_gaps,
        "citation_error_records": citation_error_records,
        "gap_rate": round(records_with_gaps / total_records, 4) if total_records else 0,
        "missing_counts": missing_counts,
        "insufficient_block_counts": insufficient_block_counts,
        "lessons": lessons,
        "recent_records": recent_records,
    }


# ---- start ----


# ---- 知识库 Chunk 管理 ----

@app.get("/api/admin/chunks")
async def admin_chunks(username: str = Depends(require_admin_or_reviewer)):
    """浏览知识库 chunks"""
    from search_engine import get_collection
    col = get_collection()
    if col.count() == 0:
        return {"chunks": [], "total": 0}
    result = col.get(include=["documents", "metadatas"], limit=200)
    chunks = []
    for i in range(min(len(result["ids"]), len(result.get("documents", [])))):
        meta = result["metadatas"][i] if result.get("metadatas") and i < len(result["metadatas"]) else {}
        chunks.append({
            "id": result["ids"][i],
            "text": (t := result["documents"][i])[:200] + ("..." if len(t) > 200 else ""),
            "lesson": meta.get("lesson", ""),
            "grade": meta.get("grade", ""),
            "chunk_type": meta.get("chunk_type", ""),
            "doc_type": meta.get("doc_type", ""),
            "source_role": meta.get("source_role", ""),
        })
    return {"chunks": chunks, "total": col.count()}


@app.post("/api/admin/chunks/delete")
async def admin_delete_chunk(req: dict, username: str = Depends(require_admin)):
    """删除单个 chunk"""
    chunk_id = req.get("chunk_id", "")
    if not chunk_id:
        raise HTTPException(status_code=400, detail="请提供 chunk_id")
    from search_engine import get_collection
    col = get_collection()
    col.delete(ids=[chunk_id])
    from search_engine import refresh_index
    refresh_index()
    audit_log(username, "admin", "delete_chunk", chunk_id)
    return {"ok": True, "message": "已删除并刷新索引"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
